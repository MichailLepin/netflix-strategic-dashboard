"""Generate the Netflix Strategic Dashboard project report as a Word document.

Usage: python generate_report.py
Output: Lepin_Novikov.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_table(doc, headers, rows):
    """Add a formatted table to the document."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_text

    return table


def generate_report():
    """Generate the full project report document."""
    doc = Document()

    # --- Title ---
    title = doc.add_heading(
        "Netflix Strategic Dashboard — Project Report", level=0
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    authors = doc.add_paragraph("Mikhail I. Lepin & Ilya R. Novikov")
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in authors.runs:
        run.font.size = Pt(14)

    program = doc.add_paragraph(
        'Master\'s Program "Business Analytics and Big Data Systems"\n'
        "March 2026"
    )
    program.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in program.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(128, 128, 128)

    # =========================================================================
    # Section 1: Business Requirements
    # =========================================================================
    doc.add_heading("1. Business Requirements", level=1)

    doc.add_heading("Business Problem", level=2)
    doc.add_paragraph(
        "Netflix operates a global streaming platform serving 280M+ subscribers "
        "across 190 countries with 17,000+ titles. The core business challenge is "
        "content overload: without effective guidance, users disengage and cancel "
        "subscriptions. The platform needs to continuously monitor user engagement, "
        "recommendation engine effectiveness, content performance, and subscription "
        "health to make data-driven decisions about content investment and user "
        "retention strategies."
    )
    doc.add_paragraph(
        "Currently, there is no consolidated single-screen view that empowers "
        "leadership to diagnose emerging trends and act quickly. This dashboard "
        "addresses that gap by providing real-time, at-a-glance intelligence "
        "for strategic decision-making."
    )

    doc.add_heading("Business Goals", level=2)
    goals = [
        "Monitor key engagement metrics (watch time, content completion) to evaluate content strategy effectiveness",
        "Track recommendation engine performance (click-through rates by algorithm type) to optimize personalization",
        "Identify churn risks across subscription tiers to prioritize retention interventions",
        "Analyze viewing patterns by device type and content category to inform platform investment",
    ]
    for g in goals:
        doc.add_paragraph(g, style="List Bullet")

    doc.add_heading("Key Questions the Dashboard Answers", level=2)
    questions = [
        "Is user engagement (completion rate, session time) improving over time? Are there seasonal patterns to account for?",
        "How effective are different recommendation algorithms (personalized vs trending vs genre-based) at driving clicks?",
        "Which subscription segments (Basic, Standard, Premium, Premium+) show the highest churn risk?",
        "How do viewing engagement patterns differ across devices (Smart TV vs Mobile vs Desktop)?",
    ]
    for q in questions:
        doc.add_paragraph(q, style="List Bullet")

    doc.add_heading("Target Audience", level=2)
    doc.add_paragraph(
        "VP of Product / Head of Content Strategy at Netflix. The dashboard is "
        "designed for executive-level decision-makers who need to monitor platform "
        "health across multiple dimensions without drilling into raw data."
    )

    # =========================================================================
    # Section 2: Data Preparation Notes
    # =========================================================================
    doc.add_heading("2. Data Preparation Notes", level=1)

    doc.add_heading("Dataset Source", level=2)
    doc.add_paragraph(
        'The base dataset is "Netflix 2025 User Behavior Dataset" from Kaggle '
        "(by Sayeeduddin, 210,000+ records across 6 interconnected tables). "
        "The dataset covers user demographics, content metadata, viewing sessions, "
        "recommendation logs, search queries, and reviews for a Netflix-style "
        "streaming platform operating in USA and Canada over 2024-2025."
    )

    doc.add_heading("Data Enrichment", level=2)
    doc.add_paragraph(
        "During initial exploratory data analysis, we discovered that the original "
        "Kaggle dataset was generated with uniformly random distributions, meaning "
        "all metrics (churn rate, watch time, recommendation CTR) were nearly "
        "identical across all segments (~15% churn for all plans, ~1.1 hrs watch "
        "time for all devices, ~15% CTR for all recommendation types). This made "
        "the data unsuitable for a meaningful business intelligence dashboard."
    )
    doc.add_paragraph(
        "To address this, we enriched the dataset by injecting realistic business "
        "patterns using an AI-assisted data augmentation script (fix_data.py). "
        "The enrichment prompt was designed based on industry knowledge about "
        "streaming platform dynamics. The following patterns were introduced:"
    )
    enrichment_items = [
        "Churn correlation with subscription tier: Basic ~14%, Standard ~7%, "
        "Premium ~4%, Premium+ ~3% — reflecting that higher-value plans retain better",
        "Device-based viewing patterns: Smart TV sessions 2.57 hrs avg vs "
        "Mobile 0.85 hrs — big screen enables lean-back, longer sessions",
        "Seasonal engagement spikes: December/January +30-50% (holiday binge), "
        "July/August +5-25% (summer), February -10-25% (post-holiday dip)",
        "Recommendation effectiveness by type: Personalized ~37% CTR vs "
        "Similar Users ~13% CTR — better algorithms produce better results",
        "Platform growth trend: +25% engagement growth over the 24-month period, "
        "simulating a maturing platform with improving content catalog",
        "Subscription tier engagement: Premium/Premium+ users watch 30% more per session, "
        "Basic users watch 30% less — reflecting plan value correlation",
        "Content type completion rates: Movies ~70% completion vs Documentaries ~12% "
        "— different content formats have different consumption patterns",
        "Genre engagement: Thriller/Sci-Fi/Drama have 25% higher session time vs "
        "Documentary/Sport/History — reflecting audience engagement differences",
    ]
    for item in enrichment_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Dataset Structure", level=2)
    table_headers = ["Table", "Rows", "Key Columns"]
    table_rows = [
        ("users.csv", "10,000", "user_id, subscription_plan, is_active, age, country, primary_device, monthly_spend"),
        ("movies.csv", "1,000", "movie_id, content_type, genre_primary, imdb_rating, duration_minutes, is_netflix_original"),
        ("watch_history.csv", "100,000", "user_id, movie_id, watch_date, device_type, watch_duration_minutes, progress_percentage"),
        ("recommendation_logs.csv", "50,000", "user_id, movie_id, recommendation_type, was_clicked, position_in_list, algorithm_version"),
        ("search_logs.csv", "25,000", "user_id, query, timestamp, search_results_count, clicked_result"),
        ("reviews.csv", "15,000", "user_id, movie_id, review_text, rating, sentiment"),
    ]
    add_table(doc, table_headers, table_rows)
    doc.add_paragraph("")
    doc.add_paragraph("Total: 201,000 records across 6 tables.")

    doc.add_heading("Data Cleaning Steps", level=2)
    cleaning_steps = [
        "Removed duplicate records by primary keys (user_id, session_id, movie_id)",
        "Dropped rows with null join keys (user_id, movie_id) to ensure referential integrity",
        "Filled missing categorical fields (subscription_plan, country, device_type) with mode values",
        "Filled missing numeric fields (age, monthly_spend) with median values",
        "Capped age outliers to 13-100 range (Netflix minimum age policy)",
        "Cast low-cardinality string columns to pandas category dtype for memory optimization",
        "Parsed date columns (watch_date, subscription_start_date) to datetime64",
        "Validated join key cardinality: 100% match rate on user_id and movie_id across all tables",
        "Total memory after optimization: 73.3 MB (within Streamlit Cloud 1GB limit)",
    ]
    for step in cleaning_steps:
        doc.add_paragraph(step, style="List Bullet")

    doc.add_heading("Key Variable Summary", level=2)
    var_headers = ["Variable", "Type", "Description"]
    var_rows = [
        ("user_id", "int64", "Unique user identifier; primary join key across all tables"),
        ("watch_duration_minutes", "float64", "Duration of a single viewing session (5-300 min)"),
        ("progress_percentage", "float64", "% of content watched in a session (0-100)"),
        ("genre_primary", "category", "Primary genre (20 categories: Adventure, Animation, Comedy, etc.)"),
        ("subscription_plan", "category", "Subscription tier: Basic, Standard, Premium, Premium+"),
        ("device_type", "category", "Viewing device: Desktop, Laptop, Mobile, Smart TV, Tablet"),
        ("watch_date", "datetime64", "Date of viewing session (2024-01-01 to 2025-12-31)"),
        ("content_type", "category", "Content format: Movie, TV Series, Documentary, Stand-up Comedy, Limited Series"),
        ("is_active", "bool", "Whether user account is active (False = churned)"),
        ("was_clicked", "bool", "Whether a recommendation was clicked by the user"),
        ("recommendation_type", "category", "Algorithm: personalized, genre_based, trending, new_releases, similar_users"),
        ("imdb_rating", "float64", "IMDb rating of content (0.5-10.0)"),
    ]
    add_table(doc, var_headers, var_rows)

    # =========================================================================
    # Section 3: Dashboard Screenshot
    # =========================================================================
    doc.add_heading("3. Dashboard Screenshot", level=1)

    doc.add_paragraph(
        "Screenshot of the deployed dashboard (captured from Streamlit Cloud URL):"
    )
    p = doc.add_paragraph()
    run = p.add_run(
        "[INSERT SCREENSHOT HERE — take a screenshot of the live dashboard at the "
        "URL below and paste it into this location]"
    )
    run.italic = True
    run.font.color.rgb = RGBColor(200, 0, 0)

    # =========================================================================
    # Section 4: Public Link
    # =========================================================================
    doc.add_heading("4. Public Link", level=1)

    doc.add_paragraph(
        "Dashboard URL: https://netflix-strategic-dashboard-mox8snlruuhsug8zwtrvsa.streamlit.app/"
    )
    doc.add_paragraph(
        "Repository: https://github.com/michaillepin/netflix-strategic-dashboard/blob/master/app.py"
    )

    # =========================================================================
    # Section 5: Dashboard and Alarm Description
    # =========================================================================
    doc.add_heading("5. Dashboard and Alarm Description", level=1)

    doc.add_heading("Dashboard Layout", level=2)
    doc.add_paragraph(
        "The Netflix Strategic Dashboard is a single-screen Streamlit web application "
        "with a Netflix-inspired dark theme (background #141414, accent #E50914). "
        "The layout uses Streamlit's wide mode to maximize screen real estate."
    )

    doc.add_heading("Dashboard Elements", level=2)
    elements = [
        "Header: Netflix-branded title row with red 'NETFLIX' logo text and "
        "'Netflix Strategic Dashboard' title with subtitle 'Content Strategy & "
        "User Engagement Analytics'.",

        "Sidebar Filters (4 interactive controls): Genre (multiselect, 20 genres), "
        "Subscription Type (multiselect: Basic/Standard/Premium/Premium+), "
        "Device Type (multiselect: Desktop/Laptop/Mobile/Smart TV/Tablet), "
        "and Time Period (date range picker, 2024-01-01 to 2025-12-31). "
        "All filters apply simultaneously to all KPIs and charts.",

        "KPI Scorecard Row (4 metric tiles): Avg Session Time (1.6 hrs overall), "
        "Content Completion Rate (51.5% overall), Recommendation CTR (23.3% overall), "
        "and Churn Rate (6.8% overall). Each tile shows a color-coded alarm border "
        "and delta indicator when filters are active.",

        "Chart 1 — Monthly Engagement Trend (Dual-axis Line Chart): "
        "Tracks completion rate (%) and average session time (hrs) over the "
        "24-month period. Reveals seasonal patterns: holiday spikes in Dec/Jan, "
        "summer boost in Jul/Aug, and Feb dip. Completion rate is the key ML "
        "target P(completion) from the recommendation engine case study. "
        "Answers: 'Is engagement improving over time?'",

        "Chart 2 — Churn Rate by Subscription Plan (Vertical Bar Chart): "
        "Shows churn risk across subscription tiers. Basic plan shows 13.5% churn "
        "(red — critical), while Premium+ shows only 3.0% (green — healthy). "
        "Color-coded bars make risk levels immediately visible. "
        "Answers: 'Which segments are we losing? Where to invest in retention?'",

        "Chart 3 — Avg Session Duration by Device (Horizontal Bar Chart): "
        "Compares average viewing session length across devices. Smart TV leads "
        "at 2.57 hrs (green) while Mobile is shortest at 0.85 hrs (red). "
        "Answers: 'Where is engagement strongest? Which platform UX needs investment?'",

        "Chart 4 — Recommendation CTR by Algorithm Type (Horizontal Bar Chart): "
        "Compares click-through rates across recommendation strategies. Personalized "
        "recommendations lead at 36.4% (green), followed by Genre-Based (28.1%), "
        "Trending (22.1%), New Releases (17.1%), and Similar Users (12.7%). "
        "Directly ties to the case study: 'Which algorithm produces the best predictions?'",
    ]
    for elem in elements:
        doc.add_paragraph(elem, style="List Bullet")

    doc.add_heading("Alarm Thresholds and Recommended Actions", level=2)
    doc.add_paragraph(
        "Each KPI tile displays a color-coded alarm border: green indicates "
        "healthy performance, yellow signals a warning that warrants monitoring, "
        "and red flags critical issues requiring immediate attention. "
        "These thresholds are also displayed in the dashboard sidebar."
    )

    alarm_headers = ["KPI", "Green (Healthy)", "Yellow (Warning)", "Red (Critical)", "Trigger Action"]
    alarm_rows = [
        ("Avg Session Time", "> 1.5 hrs", "1-1.5 hrs", "< 1 hr",
         "Evaluate content engagement strategies; investigate content quality"),
        ("Completion Rate", "> 60%", "40-60%", "< 40%",
         "Review content quality, pacing, and episode length; A/B test thumbnails"),
        ("Rec Click-Through", "> 20%", "10-20%", "< 10%",
         "Tune recommendation algorithm; review personalization model features"),
        ("Churn Rate", "< 15%", "15-25%", "> 25%",
         "Launch retention campaigns; review pricing and plan benefits; investigate support issues"),
    ]
    add_table(doc, alarm_headers, alarm_rows)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Example alarm scenario: When filtering by 'Basic' subscription tier, "
        "the Churn Rate KPI turns YELLOW (13.5%, within the 15-25% warning zone), "
        "while Avg Session Time drops to 1.1 hrs (also YELLOW). This signals that "
        "Basic subscribers are less engaged and at moderate churn risk. The recommended "
        "action is to investigate the Basic plan value proposition and consider "
        "targeted retention campaigns or feature upgrades for this segment."
    )
    doc.add_paragraph(
        "Another scenario: When filtering by 'Mobile' device, Avg Session Time "
        "drops to 0.85 hrs (RED alarm, < 1 hr threshold). This is expected behavior "
        "(mobile sessions are shorter), but the completion rate also drops. Action: "
        "optimize mobile UX for shorter, more satisfying viewing sessions."
    )

    doc.add_heading("Update Frequency", level=2)
    doc.add_paragraph(
        "In a production environment, this dashboard would be updated daily "
        "with fresh data from Netflix's data warehouse. The current implementation "
        "uses a static dataset covering January 2024 through December 2025."
    )

    # =========================================================================
    # Section 6: Data Enrichment Methodology
    # =========================================================================
    doc.add_heading("6. Data Enrichment Methodology", level=1)

    doc.add_paragraph(
        "The original Kaggle dataset had uniformly random distributions across all "
        "dimensions, making it impossible to derive actionable insights. We developed "
        "a structured AI-assisted enrichment approach to inject realistic streaming "
        "platform dynamics while preserving the original schema and row counts."
    )

    doc.add_heading("Enrichment Process", level=2)
    process_steps = [
        "Step 1 — Exploratory analysis of the original Kaggle data confirmed flat "
        "distributions: all KPIs were identical across segments (~15% churn for all "
        "plans, ~1.1 hrs watch time for all devices, ~15% CTR for all recommendation types).",
        "Step 2 — We designed a detailed enrichment prompt (data_enrichment_prompt.md, "
        "included in the repository) specifying realistic business patterns based on "
        "industry knowledge about streaming platforms.",
        "Step 3 — The enrichment script (fix_data.py) was generated from the prompt "
        "and applied multiplicative transformations to 3 of the 6 tables: users.csv "
        "(churn patterns), watch_history.csv (viewing patterns), and recommendation_logs.csv "
        "(algorithm effectiveness).",
        "Step 4 — Validation confirmed that enriched data produces meaningful variance "
        "across segments while maintaining the original schema, row counts, and data types.",
    ]
    for step in process_steps:
        doc.add_paragraph(step, style="List Bullet")

    doc.add_heading("Enrichment Prompt Summary", level=2)
    doc.add_paragraph(
        "The full enrichment prompt is available in the repository as "
        "data_enrichment_prompt.md. Key design principles:"
    )
    prompt_principles = [
        "All transformations are multiplicative — they shift means while preserving "
        "distribution shapes, rather than replacing values entirely",
        "Patterns are based on documented streaming industry dynamics: position bias "
        "in recommendations, device-based session length differences, seasonal "
        "viewing patterns, and subscription tier retention curves",
        "Random seed (42) ensures reproducibility — running fix_data.py on the "
        "original Kaggle data produces identical results",
        "3 tables modified (users, watch_history, recommendation_logs), 3 tables "
        "untouched (movies, search_logs, reviews) — only tables relevant to "
        "dashboard KPIs were enriched",
    ]
    for p in prompt_principles:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading("Before vs After Enrichment", level=2)
    comparison_headers = ["Metric", "Before (Kaggle Original)", "After (Enriched)"]
    comparison_rows = [
        ("Churn Rate (Basic)", "~15%", "13.5%"),
        ("Churn Rate (Premium+)", "~15%", "3.0%"),
        ("Avg Session (Smart TV)", "1.1 hrs", "2.57 hrs"),
        ("Avg Session (Mobile)", "1.1 hrs", "0.85 hrs"),
        ("Rec CTR (Personalized)", "~15%", "36.4%"),
        ("Rec CTR (Similar Users)", "~15%", "12.7%"),
        ("Completion (Movie)", "~10%", "70.5%"),
        ("Completion (Documentary)", "~10%", "12.4%"),
    ]
    add_table(doc, comparison_headers, comparison_rows)

    # =========================================================================
    # Section 7: Peer Review
    # =========================================================================
    doc.add_heading("7. Peer Review", level=1)

    doc.add_heading("Dashboard Reviewed", level=2)
    doc.add_paragraph(
        "Team: Marko Kolarevic, Dimitrije Dragicevic\n"
        "Dashboard: Global Music Performance & Genre Trends (Universal Music Group)\n"
        "Link: https://koyara.netlify.app/\n"
        "Tool: Custom web application (Netlify deployment)"
    )

    doc.add_heading("Overall Impression", level=2)
    doc.add_paragraph(
        "A well-designed strategic dashboard that effectively monitors the global "
        "music industry performance for Universal Music Group. The dashboard presents "
        "a clear business narrative: tracking revenue across genres, regions, and "
        "release types (singles vs albums) to inform content investment decisions."
    )

    doc.add_heading("Strengths", level=2)
    strengths = [
        "Professional visual design with consistent typography, organized grid layout, "
        "and a sophisticated color palette (gold accents, dark headers) that feels "
        "appropriate for a premium entertainment brand.",
        "Good variety of visualization types: line chart (genre trends over time), "
        "doughnut chart (market share), horizontal bar (regional revenue), stacked bar "
        "(single vs album by genre), and data table (top releases). This diversity "
        "makes the dashboard informative and visually engaging.",
        "5 well-chosen KPIs (Total Revenue, YoY Growth, Top Genre, Top Region, "
        "Single:Album Ratio) that directly support strategic decision-making.",
        "Comprehensive filter system (year, genre, region, release type, release name "
        "search) with 5 controls that enable deep drill-down analysis.",
        "Alarm/alert triggers are thoughtfully designed: YoY growth below 0%, "
        "regional drops exceeding 20%, genre share below 5%. These are actionable "
        "thresholds tied to real business decisions.",
    ]
    for s in strengths:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("Areas for Improvement", level=2)
    improvements = [
        "The dashboard could benefit from a time period filter (date range picker) "
        "in addition to the year dropdown, allowing analysis of specific quarters "
        "or months within a year for more granular seasonal insights.",
        "The Top Releases table, while informative, takes significant screen space. "
        "Consider making it collapsible or moving it to a secondary view to keep "
        "the main screen more focused on visual analytics.",
        "Consider adding a trend indicator (up/down arrow with percentage) next to "
        "each KPI to show direction of change at a glance, similar to sparklines "
        "used in financial dashboards.",
    ]
    for i in improvements:
        doc.add_paragraph(i, style="List Bullet")

    doc.add_heading("Conclusion", level=2)
    doc.add_paragraph(
        "Overall, this is a strong dashboard that demonstrates solid understanding "
        "of BI principles. The combination of clear business framing, appropriate "
        "visualization choices, and actionable alarm thresholds makes it a useful "
        "decision-making tool for music industry executives. The custom web deployment "
        "shows technical ambition beyond standard BI tools, and the result is polished "
        "and professional. Score: 8/10."
    )

    # --- Save ---
    output_file = "Lepin_Novikov.docx"
    doc.save(output_file)
    print(f"Generated: {output_file}")


if __name__ == "__main__":
    generate_report()
